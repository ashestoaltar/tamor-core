# services/file_parsing.py
"""
Centralized file parsing / text extraction for Tamor.

Core entrypoint:

    extract_text_from_file(full_path: str, mime_type: str, filename: str) -> dict

It returns a dict:

    {
        "text": "...",        # best-effort plain text for LLM/search
        "meta": {...},        # optional metadata (page count, sheets, etc.)
        "warnings": [...],    # human-readable warnings
        "parser": "backend"   # which strategy handled the file
    }

Design goals:
- Hide all format-specific logic (PDF, Word, Excel, HTML, plain text) behind
  a single function.
- Keep graceful fallbacks when optional libraries are missing.
- Preserve legacy placeholder messages so existing code that checks for
  placeholder prefixes keeps working.
"""

from __future__ import annotations

import os
import re
from typing import Any, Dict, List, Optional, Tuple


# ---------------------------------------------------------------------------
# Text Cleanup
# ---------------------------------------------------------------------------


def clean_extracted_text(text: str) -> str:
    """
    Clean up poorly formatted extracted text from PDFs.

    Handles common issues:
    - Standalone page number lines ("2/1224", "Page 5", just "42")
    - Repeated headers/footers
    - Broken lines (mid-sentence line breaks)
    - Excessive whitespace
    """
    if not text:
        return text

    lines = text.split("\n")
    cleaned_lines = []

    # Patterns for lines to remove entirely
    page_number_patterns = [
        r"^\s*\d+\s*/\s*\d+\s*$",  # "2/1224", "15 / 200"
        r"^\s*Page\s+\d+\s*$",  # "Page 5"
        r"^\s*-\s*\d+\s*-\s*$",  # "- 42 -"
        r"^\s*\[\s*\d+\s*\]\s*$",  # "[42]"
        r"^\s*\d+\s*$",  # Just a number (be careful - only if short)
    ]
    page_number_re = re.compile("|".join(page_number_patterns), re.IGNORECASE)

    # First pass: remove page number lines
    for line in lines:
        stripped = line.strip()

        # Skip empty lines for now (we'll handle spacing later)
        if not stripped:
            cleaned_lines.append("")
            continue

        # Skip standalone page numbers
        if page_number_re.match(stripped):
            continue

        # Skip very short lines that are just numbers (likely page numbers)
        if len(stripped) <= 4 and stripped.isdigit():
            continue

        cleaned_lines.append(line)

    # Second pass: merge broken lines
    # A line is "broken" if it doesn't end with sentence punctuation
    # and the next line starts with a lowercase letter
    merged_lines = []
    i = 0
    sentence_enders = '.!?:;"\'"\u201d\u2019'  # Include curly quotes

    while i < len(cleaned_lines):
        line = cleaned_lines[i]
        stripped = line.strip()

        if not stripped:
            # Preserve paragraph breaks (blank lines)
            merged_lines.append("")
            i += 1
            continue

        # Check if this line should be merged with the next
        while i + 1 < len(cleaned_lines):
            next_line = cleaned_lines[i + 1].strip()

            # Don't merge if next line is empty (paragraph break)
            if not next_line:
                break

            # Don't merge if current line ends with sentence punctuation
            if stripped and stripped[-1] in sentence_enders:
                break

            # Don't merge if next line starts with uppercase (likely new sentence/heading)
            if next_line and next_line[0].isupper():
                # But do merge if current line ends with common continuation patterns
                if not (stripped.endswith(",") or stripped.endswith("-") or
                        stripped.endswith("and") or stripped.endswith("or") or
                        stripped.endswith("the") or stripped.endswith("a")):
                    break

            # Merge the lines
            stripped = stripped + " " + next_line
            i += 1

        merged_lines.append(stripped)
        i += 1

    # Third pass: collapse multiple blank lines to one
    final_lines = []
    prev_blank = False
    for line in merged_lines:
        is_blank = not line.strip()
        if is_blank:
            if not prev_blank:
                final_lines.append("")
            prev_blank = True
        else:
            final_lines.append(line)
            prev_blank = False

    # Remove leading/trailing blank lines
    while final_lines and not final_lines[0].strip():
        final_lines.pop(0)
    while final_lines and not final_lines[-1].strip():
        final_lines.pop()

    return "\n".join(final_lines)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _ext_and_mime(
    full_path: str, mime_type: Optional[str], filename: Optional[str]
) -> Tuple[str, str, str]:
    """
    Normalize extension and MIME type.

    Returns (ext, mime_lower, name_for_display)
    """
    name = filename or os.path.basename(full_path)
    ext = os.path.splitext(name)[1].lower()
    mime_lower = (mime_type or "").lower()
    return ext, mime_lower, name


def _read_small_text_file(full_path: str) -> str:
    """
    Read a plain-text file as UTF-8, ignoring errors, with a size sanity check.

    We don't strictly enforce size caps here because the caller (e.g. _read_file_text)
    already truncates to a safe maximum, but we still avoid obviously huge files.
    """
    # Hard cap at ~5 MB for direct text read.
    max_bytes = 5 * 1024 * 1024
    try:
        size = os.path.getsize(full_path)
    except OSError:
        size = 0

    if size and size > max_bytes:
        return (
            "This file is very large for plain-text processing (>{:.1f} MB). "
            "You can still download and open it directly, but automatic "
            "summarization and search may be incomplete."
        ).format(max_bytes / (1024 * 1024))

    try:
        with open(full_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return "Error reading file contents."


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def _parse_pdf(full_path: str) -> Dict[str, Any]:
    """
    Best-effort PDF parsing using pypdf if available.

    Preserves the legacy placeholder messages so that other parts
    of the system can continue to detect non-parseable PDFs.
    Also records per-page character offsets so chunks can be mapped
    back to page numbers.
    """
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception:
        return {
            "text": (
                "This file is a PDF, but the PDF parser (pypdf) is not installed "
                "on the server yet. Install it with `pip install pypdf` to enable "
                "text extraction, summarization, and QA for PDFs."
            ),
            "meta": {},
            "warnings": [],
            "parser": "pdf-pypdf-missing",
        }

    try:
        with open(full_path, "rb") as f:
            reader = PdfReader(f)

            page_texts: List[str] = []
            page_offsets: List[int] = []
            running = 0

            for page in reader.pages:
                try:
                    t = page.extract_text() or ""
                except Exception:
                    t = ""

                # Record where this page starts in the concatenated text
                page_offsets.append(running)
                page_texts.append(t)

                # When we later do "\n".join(page_texts) we add 1 char newline
                running += len(t) + 1

        raw_text = "\n".join(page_texts).strip()
        if not raw_text:
            return {
                "text": (
                    "This PDF appears to have no extractable text (it may be a scan). "
                    "You can still open it directly from the file list."
                ),
                "meta": {"page_count": len(getattr(reader, "pages", []))},
                "warnings": [
                    "No text could be extracted from this PDF. "
                    "It may be a scanned document or image-only file."
                ],
                "parser": "pdf-pypdf2-empty",
            }

        # Clean up the extracted text (remove page numbers, merge broken lines)
        text = clean_extracted_text(raw_text)

        meta: Dict[str, Any] = {
            "page_count": len(getattr(reader, "pages", [])),
            # Note: page_offsets are based on raw text, may not match cleaned text
            "page_offsets": page_offsets,
        }
        return {
            "text": text,
            "meta": meta,
            "warnings": [],
            "parser": "pdf-pypdf2-cleaned",
        }
    except Exception as e:
        return {
            "text": f"Error extracting text from PDF: {e}",
            "meta": {},
            "warnings": [f"Exception while parsing PDF: {e}"],
            "parser": "pdf-pypdf2-error",
        }



# ---------------------------------------------------------------------------
# DOCX / Word
# ---------------------------------------------------------------------------


def _parse_docx(full_path: str) -> Dict[str, Any]:
    """
    Parse a modern Word document (.docx) using python-docx if available.
    """
    try:
        import docx  # type: ignore
    except Exception:
        return {
            "text": (
                "This file is a Word (.docx) document, but python-docx is not "
                "installed on the server yet. Install it with "
                "`pip install python-docx` to enable text extraction."
            ),
            "meta": {},
            "warnings": [],
            "parser": "docx-missing-python-docx",
        }

    try:
        d = docx.Document(full_path)
        parts: List[str] = []
        para_count = 0

        # Paragraphs
        for p in d.paragraphs:
            text = (p.text or "").strip()
            if text:
                parts.append(text)
                para_count += 1

        # Tables – flatten to simple text rows
        table_rows = 0
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(c for c in cells if c))
                    table_rows += 1

        meta: Dict[str, Any] = {
            "paragraph_count": para_count,
            "table_rows": table_rows,
        }

        text = "\n".join(parts).strip()
        if not text:
            text = (
                "This .docx file was parsed but no textual content was found. "
                "It may be mostly images or empty."
            )
            warnings = ["No textual content found in the .docx document."]
            parser = "docx-empty"
        else:
            warnings = []
            parser = "docx-python-docx"

        return {
            "text": text,
            "meta": meta,
            "warnings": warnings,
            "parser": parser,
        }
    except Exception as e:
        return {
            "text": f"Error extracting text from .docx: {e}",
            "meta": {},
            "warnings": [f".docx parsing raised an exception: {e}"],
            "parser": "docx-error",
        }


# ---------------------------------------------------------------------------
# Excel (XLSX / XLSM)
# ---------------------------------------------------------------------------


def _parse_xlsx(full_path: str, ext: str) -> Dict[str, Any]:
    """
    Parse modern Excel workbooks using openpyxl if available.

    We return a flattened text representation of visible cells, prefixed
    with sheet names, and also some structural metadata.
    """
    try:
        from openpyxl import load_workbook  # type: ignore
    except Exception:
        return {
            "text": (
                "This file is an Excel workbook (.xlsx/.xlsm), but openpyxl is not "
                "installed on the server yet. Install it with "
                "`pip install openpyxl` to enable text extraction."
            ),
            "meta": {},
            "warnings": [],
            "parser": "excel-missing-openpyxl",
        }

    try:
        wb = load_workbook(full_path, read_only=True, data_only=True)
        lines: List[str] = []
        sheet_metas: List[Dict[str, Any]] = []

        # Conservative caps so we don't explode on huge sheets.
        MAX_ROWS = 200
        MAX_COLS = 30

        truncated = False

        for ws in wb.worksheets:
            max_row = ws.max_row or 0
            max_col = ws.max_column or 0
            sheet_meta = {
                "name": ws.title,
                "max_row": max_row,
                "max_col": max_col,
            }
            sheet_metas.append(sheet_meta)

            rows_to_scan = min(max_row, MAX_ROWS) if max_row else 0
            cols_to_scan = min(max_col, MAX_COLS) if max_col else 0

            if max_row > MAX_ROWS or max_col > MAX_COLS:
                truncated = True

            if rows_to_scan == 0 or cols_to_scan == 0:
                continue

            for row in ws.iter_rows(
                min_row=1, max_row=rows_to_scan, max_col=cols_to_scan, values_only=True
            ):
                cells = [c for c in row if c not in (None, "")]
                if not cells:
                    continue
                cell_strs = [str(c) for c in cells]
                lines.append(f"{ws.title}: " + " | ".join(cell_strs))

        meta: Dict[str, Any] = {"sheets": sheet_metas}
        warnings: List[str] = []
        if truncated:
            warnings.append(
                "Excel parsing was truncated to the first "
                f"{MAX_ROWS} rows and {MAX_COLS} columns per sheet."
            )

        text = "\n".join(lines).strip()
        if not text:
            text = (
                "This Excel workbook was parsed but no textual cell data was found "
                "in the first few hundred rows/columns."
            )
            if not warnings:
                warnings.append("No non-empty cells found in initial scan.")

        return {
            "text": text,
            "meta": meta,
            "warnings": warnings,
            "parser": "excel-openpyxl",
        }
    except Exception as e:
        return {
            "text": f"Error extracting text from Excel: {e}",
            "meta": {},
            "warnings": [f"Excel parsing raised an exception: {e}"],
            "parser": "excel-error",
        }


# ---------------------------------------------------------------------------
# Legacy .xls (optional)
# ---------------------------------------------------------------------------


def _parse_xls_legacy(full_path: str) -> Dict[str, Any]:
    """
    Optional support for legacy .xls using xlrd if available.

    If xlrd is not installed, we fall back to a helpful message.
    """
    try:
        import xlrd  # type: ignore
    except Exception:
        return {
            "text": (
                "This file is an older Excel workbook (.xls). To enable parsing, "
                "install `xlrd` and ensure it is configured for .xls support."
            ),
            "meta": {},
            "warnings": [],
            "parser": "xls-missing-xlrd",
        }

    try:
        book = xlrd.open_workbook(full_path)
        lines: List[str] = []
        sheet_metas: List[Dict[str, Any]] = []

        MAX_ROWS = 200
        MAX_COLS = 30
        truncated = False

        for sheet in book.sheets():
            nrows, ncols = sheet.nrows, sheet.ncols
            sheet_metas.append(
                {"name": sheet.name, "max_row": nrows, "max_col": ncols}
            )

            rows_to_scan = min(nrows, MAX_ROWS)
            cols_to_scan = min(ncols, MAX_COLS)

            if nrows > MAX_ROWS or ncols > MAX_COLS:
                truncated = True

            for r in range(rows_to_scan):
                cells = [sheet.cell_value(r, c) for c in range(cols_to_scan)]
                cells = [c for c in cells if c not in (None, "")]
                if not cells:
                    continue
                cell_strs = [str(c) for c in cells]
                lines.append(f"{sheet.name}: " + " | ".join(cell_strs))

        meta: Dict[str, Any] = {"sheets": sheet_metas}
        warnings: List[str] = []
        if truncated:
            warnings.append(
                "Excel (.xls) parsing was truncated to the first "
                f"{MAX_ROWS} rows and {MAX_COLS} columns per sheet."
            )

        text = "\n".join(lines).strip()
        if not text:
            text = (
                "This Excel (.xls) workbook was parsed but no textual cell data "
                "was found in the scanned range."
            )

        return {
            "text": text,
            "meta": meta,
            "warnings": warnings,
            "parser": "xls-xlrd",
        }
    except Exception as e:
        return {
            "text": f"Error extracting text from Excel (.xls): {e}",
            "meta": {},
            "warnings": [f".xls parsing raised an exception: {e}"],
            "parser": "xls-error",
        }


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------


def _parse_html(full_path: str) -> Dict[str, Any]:
    """
    Very lightweight HTML parsing using the standard library.

    We strip tags, keep a bit of structure (headings, paragraphs, list items),
    and otherwise just return the visible text.
    """
    from html.parser import HTMLParser

    class _HTMLTextExtractor(HTMLParser):
        def __init__(self) -> None:
            super().__init__()
            self.parts: List[str] = []
            self._in_title = False

        def handle_starttag(self, tag, attrs):
            t = tag.lower()
            if t in ("p", "br", "li", "h1", "h2", "h3", "h4"):
                self.parts.append("\n")
            if t == "title":
                self._in_title = True

        def handle_endtag(self, tag):
            if tag.lower() == "title":
                self._in_title = False

        def handle_data(self, data):
            if not data:
                return
            text = data.strip()
            if not text:
                return
            self.parts.append(text + " ")

    try:
        with open(full_path, "r", encoding="utf-8", errors="ignore") as f:
            html = f.read()
    except Exception:
        return {
            "text": "Error reading HTML file contents.",
            "meta": {},
            "warnings": ["Failed to read HTML file from disk."],
            "parser": "html-error-read",
        }

    parser = _HTMLTextExtractor()
    try:
        parser.feed(html)
        text = "".join(parser.parts)
        text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    except Exception as e:
        return {
            "text": f"Error extracting text from HTML: {e}",
            "meta": {},
            "warnings": [f"HTML parsing raised an exception: {e}"],
            "parser": "html-error-parse",
        }

    if not text:
        text = (
            "This HTML file was parsed but no meaningful textual content was found."
        )

    return {
        "text": text,
        "meta": {},
        "warnings": [],
        "parser": "html-stdlib",
    }


# ---------------------------------------------------------------------------
# EPUB
# ---------------------------------------------------------------------------


def _parse_epub(full_path: str) -> Dict[str, Any]:
    """
    Parse an EPUB ebook and extract text content.

    EPUBs contain XHTML documents. We extract text from each,
    preserving chapter structure in metadata.
    """
    try:
        import ebooklib
        from ebooklib import epub
    except ImportError:
        return {
            "text": (
                "This file is an EPUB ebook, but ebooklib is not installed. "
                "Install it with `pip install ebooklib` to enable text extraction."
            ),
            "meta": {},
            "warnings": [],
            "parser": "epub-missing-ebooklib",
        }

    try:
        book = epub.read_epub(full_path)

        # Extract metadata
        title = book.get_metadata('DC', 'title')
        title = title[0][0] if title else None

        author = book.get_metadata('DC', 'creator')
        author = author[0][0] if author else None

        # Extract text from each document
        chapters = []
        full_text_parts = []
        chapter_offsets = []  # Track where each chapter starts in full text
        current_offset = 0

        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                content = item.get_content().decode('utf-8', errors='ignore')

                # Strip HTML tags (simple approach)
                from html.parser import HTMLParser

                class TextExtractor(HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self.text_parts = []

                    def handle_data(self, data):
                        text = data.strip()
                        if text:
                            self.text_parts.append(text)

                extractor = TextExtractor()
                try:
                    extractor.feed(content)
                    chapter_text = ' '.join(extractor.text_parts)
                except Exception:
                    chapter_text = ""

                if chapter_text.strip():
                    # Record chapter offset
                    chapter_offsets.append(current_offset)

                    # Try to get chapter title from first line
                    lines = chapter_text.strip().split('\n')
                    chapter_title = lines[0][:100] if lines else f"Chapter {len(chapters) + 1}"

                    chapters.append({
                        "index": len(chapters),
                        "title": chapter_title,
                        "char_count": len(chapter_text),
                    })

                    full_text_parts.append(chapter_text)
                    current_offset += len(chapter_text) + 2  # +2 for \n\n separator

        full_text = '\n\n'.join(full_text_parts)

        if not full_text.strip():
            return {
                "text": "This EPUB was parsed but no text content was found.",
                "meta": {"title": title, "author": author},
                "warnings": ["No extractable text in EPUB"],
                "parser": "epub-empty",
            }

        meta = {
            "title": title,
            "author": author,
            "chapter_count": len(chapters),
            "chapters": chapters,
            "chapter_offsets": chapter_offsets,
        }

        return {
            "text": full_text,
            "meta": meta,
            "warnings": [],
            "parser": "epub-ebooklib",
        }

    except Exception as e:
        return {
            "text": f"Error extracting text from EPUB: {e}",
            "meta": {},
            "warnings": [f"EPUB parsing failed: {e}"],
            "parser": "epub-error",
        }


# ---------------------------------------------------------------------------
# Fallback / dispatcher
# ---------------------------------------------------------------------------


def extract_text_from_file(
    full_path: str, mime_type: Optional[str], filename: Optional[str] = None
) -> Dict[str, Any]:
    """
    Main entrypoint for file parsing.

    Given a path, MIME type, and filename, return a dict with:
    - text
    - meta
    - warnings
    - parser

    This function is intentionally conservative: if it cannot confidently
    parse a file, it returns one of the legacy placeholder messages so that
    higher-level code can detect and skip it for semantic indexing.
    """
    ext, mime_lower, name = _ext_and_mime(full_path, mime_type, filename or "")

    # Route by type
    # PDFs
    if mime_lower == "application/pdf" or ext == ".pdf":
        return _parse_pdf(full_path)

    # Word / DOCX
    if ext == ".docx" or mime_lower == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _parse_docx(full_path)

    # Excel – modern formats
    if ext in (".xlsx", ".xlsm") or (
        mime_lower
        in (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel.sheet.macroenabled.12",
        )
    ):
        return _parse_xlsx(full_path, ext)

    # Legacy .xls
    if ext == ".xls" or mime_lower == "application/vnd.ms-excel":
        return _parse_xls_legacy(full_path)

    # HTML
    if ext in (".html", ".htm") or mime_lower == "text/html":
        return _parse_html(full_path)

    # EPUB ebooks
    if ext == ".epub" or mime_lower == "application/epub+zip":
        return _parse_epub(full_path)

    # Text-like / code / markdown etc.
    text_like_exts = {
        ".txt",
        ".md",
        ".markdown",
        ".json",
        ".js",
        ".ts",
        ".jsx",
        ".tsx",
        ".py",
        ".lisp",
        ".html",
        ".css",
        ".csv",
        ".yml",
        ".yaml",
        ".ini",
        ".cfg",
        ".conf",
        ".toml",
        ".xml",
        ".sh",
        ".bat",
        ".ps1",
        ".c",
        ".cpp",
        ".h",
        ".hpp",
        ".java",
        ".cs",
        ".rs",
    }

    if mime_lower.startswith("text/") or ext in text_like_exts:
        text = _read_small_text_file(full_path)
        return {
            "text": text,
            "meta": {},
            "warnings": [],
            "parser": "plain-text",
        }

    # Fallback: treat unknown types as non-plain-text for now.
    return {
        "text": (
            "This file is not a plain-text type. "
            "You can still download and open it, but automated summarization "
            "and search will be limited until a richer parser is added."
        ),
        "meta": {},
        "warnings": [],
        "parser": "unsupported-binary",
    }
