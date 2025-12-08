# api/services/structured_parsing.py

"""
Helpers for extracting lightweight structure from files:
- PDFs: pages + inferred headings
- DOCX: headings + table shapes
- XLSX: sheets + header row + row counts

Phase 2.6: these are intentionally simple and can be improved later, but
they are "real" enough to make summaries page/sheet/heading aware.
"""

from __future__ import annotations

from typing import Any, Dict, List, Optional, Tuple
import os


def extract_pdf_structure(file_path: str) -> Dict[str, Any]:
    """
    Return a dict like:
    {
      "type": "pdf",
      "pages": [
        {"index": 1, "heading": "Intro", "char_count": 1234},
        {"index": 2, "heading": "Installation", "char_count": 2048},
      ]
    }

    Implementation details:
    - Use PyPDF2 if available.
    - For each page, record:
      - index (1-based)
      - char_count
      - a naive heading: first non-empty line, truncated to 120 chars.
    """
    try:
        from PyPDF2 import PdfReader  # type: ignore
    except Exception:
        # No PDF parser installed → we can still say "it's a PDF"
        return {
            "type": "pdf",
            "pages": [],
        }

    pages: List[Dict[str, Any]] = []
    try:
        with open(file_path, "rb") as f:
            reader = PdfReader(f)
            for idx, page in enumerate(reader.pages, start=1):
                try:
                    text = page.extract_text() or ""
                except Exception:
                    text = ""

                char_count = len(text)
                heading = ""

                if text:
                    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
                    if lines:
                        heading = lines[0][:120]

                pages.append(
                    {
                        "index": idx,
                        "heading": heading,
                        "char_count": char_count,
                    }
                )
    except Exception:
        # If reading fails entirely, just return an empty structure
        return {
            "type": "pdf",
            "pages": [],
        }

    return {
        "type": "pdf",
        "pages": pages,
    }


def extract_docx_structure(file_path: str) -> Dict[str, Any]:
    """
    Return a dict like:
    {
      "type": "docx",
      "headings": [
        {"level": 1, "text": "Overview"},
        {"level": 2, "text": "Safety Warnings"},
      ],
      "tables": [
        {"index": 0, "approx_rows": 10, "approx_cols": 4}
      ]
    }

    Implementation details:
    - Headings: paragraphs whose style name starts with "Heading".
      We try to parse the level (e.g., "Heading 1", "Heading 2").
    - Tables: we just record number of rows and columns per table.
    """
    try:
        import docx  # type: ignore
    except Exception:
        return {
            "type": "docx",
            "headings": [],
            "tables": [],
        }

    headings: List[Dict[str, Any]] = []
    tables: List[Dict[str, Any]] = []

    try:
        d = docx.Document(file_path)

        # Headings: paragraphs whose style name starts with "Heading"
        for p in d.paragraphs:
            text = (p.text or "").strip()
            if not text:
                continue

            style = getattr(p, "style", None)
            style_name = getattr(style, "name", "") if style is not None else ""
            style_name_lower = style_name.lower()

            if not style_name_lower.startswith("heading"):
                continue

            # Try to extract a level (e.g., "Heading 1", "Heading 2")
            level: Optional[int] = None
            parts = style_name.split()
            for part in reversed(parts):
                if part.isdigit():
                    level = int(part)
                    break

            headings.append(
                {
                    "level": level or 1,
                    "text": text[:200],
                }
            )

        # Tables: just record rough shape
        for idx, tbl in enumerate(d.tables):
            rows = len(tbl.rows)
            cols = len(tbl.columns)
            tables.append(
                {
                    "index": idx,
                    "approx_rows": rows,
                    "approx_cols": cols,
                }
            )

    except Exception:
        return {
            "type": "docx",
            "headings": [],
            "tables": [],
        }

    return {
        "type": "docx",
        "headings": headings,
        "tables": tables,
    }


def extract_excel_structure(file_path: str) -> Dict[str, Any]:
    """
    Return a dict like:
    {
      "type": "xlsx",
      "sheets": [
        {
          "name": "BOM",
          "headers": ["PartNumber", "Description", "LengthMM"],
          "row_count": 123,
        },
        ...
      ]
    }

    Implementation details:
    - We use openpyxl in read_only mode.
    - For each sheet, we:
      - record the sheet name,
      - try to infer a header row by scanning the first few rows until
        we hit a non-blank row,
      - record total row_count from ws.max_row.
    """
    try:
        from openpyxl import load_workbook  # type: ignore
    except Exception:
        return {
            "type": "xlsx",
            "sheets": [],
        }

    sheets: List[Dict[str, Any]] = []

    try:
        wb = load_workbook(file_path, read_only=True, data_only=True)
        MAX_HEADER_SCAN_ROWS = 10

        for ws in wb.worksheets:
            max_row = ws.max_row or 0

            headers: List[str] = []
            # Scan first few rows to guess a header row
            scan_rows = min(MAX_HEADER_SCAN_ROWS, max_row) if max_row else 0
            if scan_rows > 0:
                for row in ws.iter_rows(
                    min_row=1, max_row=scan_rows, values_only=True
                ):
                    values = list(row)
                    if not any(v not in (None, "") for v in values):
                        continue
                    headers = [
                        str(v).strip()
                        for v in values
                        if v not in (None, "")
                    ]
                    break

            sheets.append(
                {
                    "name": ws.title,
                    "headers": headers,
                    "row_count": max_row,
                }
            )

    except Exception:
        return {
            "type": "xlsx",
            "sheets": [],
        }

    return {
        "type": "xlsx",
        "sheets": sheets,
    }


def extract_structure_for_mime(
    mime_type: str, file_path: str
) -> Optional[Dict[str, Any]]:
    """
    Convenience entrypoint: choose appropriate extractor based on MIME type
    or file extension. Returns None if we don’t have a structure extractor
    for this file type.
    """
    mt = (mime_type or "").lower()
    _, ext = os.path.splitext(file_path.lower())

    # PDF
    if mt == "application/pdf" or ext == ".pdf":
        return extract_pdf_structure(file_path)

    # Word / DOCX
    if mt in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or ext in (".docx", ".doc"):
        return extract_docx_structure(file_path)

    # Excel
    if mt in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
        "application/vnd.ms-excel.sheet.macroenabled.12",
    ) or ext in (".xlsx", ".xls", ".xlsm"):
        return extract_excel_structure(file_path)

    return None

